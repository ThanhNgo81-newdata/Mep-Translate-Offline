# -*- coding: utf-8 -*-
"""
MEP Translator Offline - Streamlit + local Hugging Face model.

Runtime does NOT use Anthropic/OpenAI/Google API.
The model is loaded only from ./models/nllb-200-distilled-600M after
download_models.py has been run once on a machine with Internet access.

Run:
    streamlit run mep_translator_offline.py --server.address 0.0.0.0 --server.port 8501
"""
import io
import re
from copy import copy
from pathlib import Path

import streamlit as st
import torch
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
import openpyxl
import fitz
from PIL import Image

try:
    import pytesseract
except ImportError:
    pytesseract = None


BASE_DIR = Path(__file__).resolve().parent
MODEL_DIR = BASE_DIR / "models" / "nllb-200-distilled-600M"
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"

LANGUAGES = {
    "eng_Latn": "Tiếng Anh",
    "vie_Latn": "Tiếng Việt",
    "zho_Hans": "Tiếng Trung giản thể",
    "zho_Hant": "Tiếng Trung phồn thể",
    "jpn_Jpan": "Tiếng Nhật",
    "kor_Hang": "Tiếng Hàn",
    "fra_Latn": "Tiếng Pháp",
    "deu_Latn": "Tiếng Đức",
    "spa_Latn": "Tiếng Tây Ban Nha",
    "rus_Cyrl": "Tiếng Nga",
    "tha_Thai": "Tiếng Thái",
    "por_Latn": "Tiếng Bồ Đào Nha",
    "ita_Latn": "Tiếng Ý",
    "ind_Latn": "Tiếng Indonesia",
    "msa_Latn": "Tiếng Mã Lai",
    "khm_Khmr": "Tiếng Khmer",
    "lao_Laoo": "Tiếng Lào",
    "arb_Arab": "Tiếng Ả Rập",
    "hin_Deva": "Tiếng Hindi",
}

# NLLB language codes for OCR.
OCR_LANGS = {
    "eng_Latn": "eng", "vie_Latn": "vie", "zho_Hans": "chi_sim",
    "zho_Hant": "chi_tra", "jpn_Jpan": "jpn", "kor_Hang": "kor",
    "fra_Latn": "fra", "deu_Latn": "deu", "spa_Latn": "spa",
    "rus_Cyrl": "rus", "tha_Thai": "tha", "por_Latn": "por",
    "ita_Latn": "ita", "ind_Latn": "ind", "msa_Latn": "msa",
    "khm_Khmr": "khm", "lao_Laoo": "lao", "arb_Arab": "ara",
    "hin_Deva": "hin",
}

GLOSSARY = (
    "MEP terminology. Preserve equipment tags, standards, model numbers, "
    "drawing numbers and engineering units exactly. Examples: AHU, FCU, VRV, "
    "VRF, MCCB, ACB, DB, TCVN, QCVN, ASHRAE, NFPA, ASME, SMACNA, IEC, JIS, "
    "kW, CFM, m3/h, Pa, mmAq, kVA, mm, A, V, Hz."
)

# Protect common technical tokens from machine translation.
PROTECT_RE = re.compile(
    r"\b(?:AHU|FCU|VRV|VRF|MCCB|ACB|DB|BMS|VAV|FAHU|PAU|"
    r"TCVN|QCVN|ASHRAE|NFPA|ASME|SMACNA|IEC|JIS|ISO|EN|"
    r"\d+(?:\.\d+)?\s?(?:kW|MW|W|kVA|V|A|Hz|Pa|kPa|bar|CFM|m3/h|mmAq|mm|m|°C|C))"
    r"(?:[-_/A-Za-z0-9.]*)\b",
    re.IGNORECASE,
)

st.set_page_config(page_title="MEP Translator Offline", page_icon="🛠️", layout="wide")


@st.cache_resource(show_spinner="Đang nạp mô hình dịch cục bộ…")
def load_model():
    if not MODEL_DIR.exists():
        raise FileNotFoundError(
            f"Chưa có mô hình tại {MODEL_DIR}. Hãy chạy download_models.py trước."
        )
    tokenizer = AutoTokenizer.from_pretrained(
        MODEL_DIR, local_files_only=True, use_fast=False
    )
    model = AutoModelForSeq2SeqLM.from_pretrained(
        MODEL_DIR, local_files_only=True
    )
    model.to(DEVICE)
    model.eval()
    return tokenizer, model


def protect_tokens(text: str):
    saved = []

    def repl(m):
        key = f"ZXQTECH{len(saved)}QXZ"
        saved.append((key, m.group(0)))
        return key

    return PROTECT_RE.sub(repl, text), saved


def restore_tokens(text: str, saved):
    for key, value in saved:
        text = text.replace(key, value)
        text = text.replace(key.lower(), value)
    return text


def translate_batch(texts, src, tgt, batch_size=8, max_length=512):
    if not texts:
        return []

    tokenizer, model = load_model()
    results = []

    for start in range(0, len(texts), batch_size):
        original_batch = texts[start:start + batch_size]
        protected = []
        maps = []
        for t in original_batch:
            p, mapping = protect_tokens(t)
            protected.append(p)
            maps.append(mapping)

        tokenizer.src_lang = src
        inputs = tokenizer(
            protected,
            return_tensors="pt",
            padding=True,
            truncation=True,
            max_length=max_length,
        )
        inputs = {k: v.to(DEVICE) for k, v in inputs.items()}

        with torch.inference_mode():
            generated = model.generate(
                **inputs,
                forced_bos_token_id=tokenizer.convert_tokens_to_ids(tgt),
                max_length=max_length,
                num_beams=4,
                early_stopping=True,
            )

        decoded = tokenizer.batch_decode(generated, skip_special_tokens=True)
        results.extend(
            restore_tokens(x.strip(), mapping)
            for x, mapping in zip(decoded, maps)
        )

    return results


def translate_text(text, src, tgt):
    if not text or not text.strip() or src == tgt:
        return text
    return translate_batch([text], src, tgt, batch_size=1)[0]


def add_paragraph_after(paragraph, text, italic=False):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    run.italic = italic
    return new_para


def translate_docx(data, src, tgt, bilingual, progress):
    doc = Document(io.BytesIO(data))
    targets = [p for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                targets.extend(p for p in cell.paragraphs if p.text.strip())

    texts = [p.text.strip() for p in targets]
    translations = translate_batch(texts, src, tgt)

    for i, (p, tr) in enumerate(zip(targets, translations), 1):
        if bilingual:
            add_paragraph_after(p, tr, italic=True)
        else:
            # Keep the paragraph style/format as much as possible.
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = tr
            else:
                p.add_run(tr)
        progress(i / max(1, len(targets)))

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def copy_cell_style(src, dst):
    try:
        dst.font = copy(src.font)
        dst.fill = copy(src.fill)
        dst.border = copy(src.border)
        dst.alignment = copy(src.alignment)
        dst.number_format = src.number_format
        dst.protection = copy(src.protection)
    except Exception:
        pass

from openpyxl.cell.cell import MergedCell

def translate_xlsx(data, src, tgt, bilingual, cb):
    from openpyxl import load_workbook
    import os

    wb = load_workbook(data)
    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                text = cell.value
                if text and isinstance(text, str):
                    tr = cb(text, src, tgt, bilingual)

                    # Nếu là MergedCell thì ghi vào ô gốc của vùng merge
                    if isinstance(cell, MergedCell):
                        for merged_range in sheet.merged_cells.ranges:
                            if (cell.row, cell.column) in merged_range.cells:
                                top_left = sheet.cell(merged_range.min_row, merged_range.min_col)
                                top_left.value = tr
                                break
                    else:
                        cell.value = tr

    # Lưu file dịch ra cùng tên + _translated.xlsx
    base, ext = os.path.splitext(data)
    out_file = f"{base}_translated{ext}"
    wb.save(out_file)
    return out_file




def pdf_to_docx(data, src, tgt, bilingual, page_from, page_to, progress):
    pdf = fitz.open(stream=data, filetype="pdf")
    out_doc = Document()
    pages = range(page_from - 1, page_to)
    pages = list(pages)

    for idx, pno in enumerate(pages, 1):
        page = pdf.load_page(pno)
        blocks = page.get_text("blocks")
        blocks = sorted(blocks, key=lambda b: (round(b[1], 1), round(b[0], 1)))
        texts = [b[4].strip() for b in blocks if b[4] and b[4].strip()]
        if not texts:
            continue

        out_doc.add_heading(f"Trang {pno + 1}", level=2)
        translations = translate_batch(texts, src, tgt)

        for orig, tr in zip(texts, translations):
            if bilingual:
                p = out_doc.add_paragraph(orig)
                if p.runs:
                    p.runs[0].italic = True
                out_doc.add_paragraph(tr)
            else:
                out_doc.add_paragraph(tr)

        if idx < len(pages):
            out_doc.add_page_break()
        progress(idx / max(1, len(pages)))

    pdf.close()
    out = io.BytesIO()
    out_doc.save(out)
    return out.getvalue()


def ocr_image(data, src, tgt, bilingual):
    if pytesseract is None:
        raise RuntimeError("Chưa cài pytesseract. Cài thêm pytesseract + Tesseract OCR.")
    lang = OCR_LANGS.get(src, "eng")
    text = pytesseract.image_to_string(Image.open(io.BytesIO(data)), lang=lang)
    if not text.strip():
        return ""
    tr = translate_text(text, src, tgt)
    return f"{text}\n\n--- BẢN DỊCH ---\n\n{tr}" if bilingual else tr


# ---------------- UI ----------------
st.title("🛠️ MEP TRANSLATOR — OFFLINE")
st.caption(
    "Không dùng Anthropic/OpenAI API. Tất cả nội dung được xử lý trên máy chủ nội bộ "
    "bằng mô hình dịch cục bộ."
)

if not MODEL_DIR.exists():
    st.error("Chưa có model offline. Hãy chạy `python download_models.py` một lần trước khi chạy app.")
    st.stop()

with st.sidebar:
    st.header("Cấu hình dịch")
    src = st.selectbox(
        "Ngôn ngữ nguồn",
        list(LANGUAGES),
        format_func=lambda x: LANGUAGES[x],
        index=0,
    )
    tgt = st.selectbox(
        "Ngôn ngữ đích",
        list(LANGUAGES),
        format_func=lambda x: LANGUAGES[x],
        index=1,
    )
    bilingual = st.checkbox("🈯 Giữ bản gốc + bản dịch", value=True)
    batch_size = st.slider("Batch dịch", 1, 16, 8)
    st.caption(f"Thiết bị chạy: **{DEVICE.upper()}**")
    st.caption(GLOSSARY)

uploaded = st.file_uploader(
    "Tải tài liệu",
    type=["pdf", "docx", "xlsx", "xls", "txt", "jpg", "jpeg", "png", "webp"],
)

if uploaded:
    data = uploaded.getvalue()
    ext = uploaded.name.lower().rsplit(".", 1)[-1]
    st.write(f"**Tệp:** {uploaded.name} · **{len(data)/1024:.0f} KB**")

    page_from = page_to = None
    if ext == "pdf":
        pdf = fitz.open(stream=data, filetype="pdf")
        n = pdf.page_count
        pdf.close()
        c1, c2 = st.columns(2)
        page_from = int(c1.number_input("Từ trang", 1, n, 1))
        page_to = int(c2.number_input("Đến trang", page_from, n, n))

    if st.button("🌐 Dịch tài liệu", type="primary"):
        progress = st.progress(0, text="Đang chuẩn bị…")

        def cb(v):
            progress.progress(min(1.0, float(v)), text=f"Đang dịch… {int(v*100)}%")

        try:
            if ext == "docx":
                result = translate_docx(data, src, tgt, bilingual, cb)
                out_name = Path(uploaded.name).stem + "_translated.docx"
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

            elif ext in ("xlsx", "xls"):
                if ext == "xls":
                    st.warning("openpyxl không ghi trực tiếp XLS cũ. Hãy lưu thành XLSX rồi dịch.")
                    st.stop()
                result = translate_xlsx(data, src, tgt, bilingual, cb)
                out_name = Path(uploaded.name).stem + "_translated.xlsx"
                mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

            elif ext == "pdf":
                result = pdf_to_docx(data, src, tgt, bilingual, page_from, page_to, cb)
                out_name = Path(uploaded.name).stem + "_translated.docx"
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

            elif ext == "txt":
                text = data.decode("utf-8", errors="replace")
                lines = [x for x in text.splitlines() if x.strip()]
                trs = translate_batch(lines, src, tgt, batch_size=batch_size)
                if bilingual:
                    output = []
                    for a, b in zip(lines, trs):
                        output.extend([a, b])
                    result = "\n".join(output).encode("utf-8")
                else:
                    result = "\n".join(trs).encode("utf-8")
                out_name = Path(uploaded.name).stem + "_translated.txt"
                mime = "text/plain"

            else:
                text = ocr_image(data, src, tgt, bilingual)
                st.text_area("Kết quả OCR + dịch", text, height=400)
                result = text.encode("utf-8")
                out_name = Path(uploaded.name).stem + "_translated.txt"
                mime = "text/plain"

            progress.progress(1.0, text="Hoàn tất.")
            st.success("Đã dịch xong.")
            st.download_button(
                "⬇️ Tải file kết quả",
                data=result,
                file_name=out_name,
                mime=mime,
            )
        except Exception as e:
            st.exception(e)
